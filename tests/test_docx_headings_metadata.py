from __future__ import annotations

import importlib.util as importlib_util
from pathlib import Path

import pytest

from oscillink.ingest.chunk import chunk_paragraphs
from oscillink.ingest.extract import extract_text

pytestmark = pytest.mark.skipif(
    importlib_util.find_spec("docx") is None, reason="python-docx not installed"
)


def test_docx_headings_propagate(tmp_path: Path) -> None:
    # Create a temporary .docx with a Heading 1 and a paragraph
    from docx import Document  # type: ignore[import-not-found]

    doc_path = tmp_path / "sample.docx"
    d = Document()
    d.add_heading("Section A", level=1)
    d.add_paragraph("This is content under Section A.")
    d.save(str(doc_path))

    ex = extract_text([str(doc_path)], parser="auto")
    assert ex and ex[0].pages, "extraction produced no pages"
    pages = ex[0].pages
    ch = chunk_paragraphs(pages)
    assert ch.chunks, "no chunks produced"
    # At least one chunk should carry section metadata
    metas = [c.meta for c in ch.chunks if c.meta]
    assert metas, "missing metadata on chunks"
    has_title = any(m.get("section_title") == "Section A" for m in metas)  # type: ignore[union-attr]
    assert has_title, f"expected section_title 'Section A' in chunk metadata; got {metas}"
