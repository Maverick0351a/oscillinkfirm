from __future__ import annotations

import datetime as _dt
import hashlib
import json
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import APIRouter
from pydantic import BaseModel, Field

from oscillink.ingest.query_service import query_index

router = APIRouter()


class ReportRequest(BaseModel):
    title: str = Field("OPRA Report")
    index_path: str
    q: str = Field(..., description="Main prompt/question to address in the report")
    backend: str = Field("jsonl")
    k: int = Field(5, ge=1, le=20)
    embed_model: str = Field("bge-small-en-v1.5")
    meta_path: Optional[str] = None
    epsilon: float = Field(1e-3, ge=0.0)
    tau: float = Field(0.30, ge=0.0, le=1.0)
    e2e: bool = False
    fmt: str = Field("docx", pattern=r"^(docx|txt)$")
    out_path: Optional[str] = Field(None, description="Optional output path; will be created under /data/reports if not provided")


class ReportResponse(BaseModel):
    path: str
    format: str
    size: int
    bundle_hash: str
    receipt: Optional[dict] = None
    ingest_receipt: Optional[dict] = None
    settle_receipt: Optional[dict] = None
    citations: Optional[List[dict]] = None
    abstain: Optional[bool] = None
    reason: Optional[str] = None


def _slugify(s: str) -> str:
    keep = [c if c.isalnum() or c in ("-", "_") else "-" for c in s.strip().lower()]
    return "".join(keep).strip("-") or "report"


def _format_extractive(q: str, items: List[dict]) -> str:
    lines = [f"Q: {q}", "", "Based on the top retrieved passages:"]
    for i, it in enumerate(items[:5], 1):
        src = it.get("source_path")
        page = it.get("page_number")
        start = it.get("start")
        end = it.get("end")
        sc = it.get("score")
        short = os.path.basename(str(src)) if src else "unknown"
        lines.append(f"{i}. {short} (page {page}, lines {start}-{end}) — score={sc}")
    lines.append("")
    lines.append("Answer: The information above contains the most relevant citations. Refer to the listed sources and passages.")
    return "\n".join(lines)


def _citations_from_qres(qres: Dict[str, Any], *, use_bundle: bool) -> List[dict]:
    items: List[dict] = []
    if use_bundle and isinstance(qres.get("bundle"), list):
        for b in qres["bundle"]:
            if isinstance(b, dict):
                items.append({
                    "score": b.get("score"),
                    "source_path": b.get("source_path"),
                    "page_number": b.get("page_number"),
                    "start": b.get("start"),
                    "end": b.get("end"),
                })
    elif isinstance(qres.get("results"), list):
        items = list(qres["results"])  # shallow copy
    return items


def _bundle_hash(payload: Dict[str, Any]) -> str:
    data = json.dumps(payload, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(data).hexdigest()


def _write_txt(out_path: Path, title: str, q: str, answer: str, citations: List[dict], receipts: Dict[str, Any]) -> None:
    lines: List[str] = []
    lines.append(title)
    lines.append("=" * len(title))
    lines.append("")
    lines.append(_dt.datetime.utcnow().isoformat() + "Z")
    lines.append("")
    lines.append("Question")
    lines.append("-" * 8)
    lines.append(q)
    lines.append("")
    lines.append("Answer")
    lines.append("-" * 6)
    lines.append(answer)
    lines.append("")
    lines.append("Citations")
    lines.append("-" * 9)
    for i, c in enumerate(citations, 1):
        lines.append(f"{i}. {c.get('source_path')} page {c.get('page_number')} lines {c.get('start')}-{c.get('end')} score={c.get('score')}")
    lines.append("")
    lines.append("Receipts")
    lines.append("-" * 8)
    lines.append(json.dumps(receipts, indent=2, sort_keys=True))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")


def _write_docx(out_path: Path, title: str, q: str, answer: str, citations: List[dict], receipts: Dict[str, Any]) -> bool:
    try:
        from docx import Document  # type: ignore

        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(_dt.datetime.utcnow().isoformat() + "Z")
        doc.add_heading("Question", level=1)
        doc.add_paragraph(q)
        doc.add_heading("Answer", level=1)
        for para in answer.splitlines():
            doc.add_paragraph(para)
        doc.add_heading("Citations", level=1)
        for i, c in enumerate(citations, 1):
            line = f"{i}. {c.get('source_path')} page {c.get('page_number')} lines {c.get('start')}-{c.get('end')} score={c.get('score')}"
            doc.add_paragraph(line, style="List Number")
        doc.add_heading("Receipts", level=1)
        doc.add_paragraph(json.dumps(receipts, indent=2, sort_keys=True))
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return True
    except Exception:
        return False


@router.post("/report", response_model=ReportResponse)
def create_report(req: ReportRequest) -> Dict[str, Any]:
    # Accept variant backends like faiss:hnsw by keeping base kind
    backend_flag = req.backend.split(":", 1)[0]
    req.backend = backend_flag if backend_flag in {"jsonl", "faiss"} else "jsonl"
    qres = query_index(
        index_path=req.index_path,
        backend=req.backend,
        q=req.q,
        k=req.k,
        embed_model=req.embed_model,
        meta_path=req.meta_path,
        e2e=req.e2e,
        epsilon=req.epsilon,
        tau=req.tau,
    )
    items = _citations_from_qres(qres, use_bundle=req.e2e)
    answer = _format_extractive(req.q, items)

    receipts: Dict[str, Any] = {
        "receipt": qres.get("receipt"),
        "ingest_receipt": qres.get("ingest_receipt"),
        "settle_receipt": qres.get("settle_receipt"),
        "abstain": qres.get("abstain"),
        "reason": qres.get("reason"),
    }
    bh_payload = {
        "title": req.title,
        "q": req.q,
        "index_path": req.index_path,
        "receipts": receipts,
    }
    bundle_hash = _bundle_hash(bh_payload)

    base_dir = Path(os.getenv("OPRA_REPORTS_DIR", "/data/reports"))
    stem = _slugify(req.title) + "-" + _dt.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    if req.out_path:
        out_path = Path(req.out_path)
    else:
        out_path = base_dir / (stem + (".docx" if req.fmt == "docx" else ".txt"))

    # Generate report file
    ok_docx = False
    if req.fmt == "docx":
        ok_docx = _write_docx(out_path, req.title, req.q, answer, items, receipts)
        if not ok_docx:
            # Fallback to txt if docx fails
            out_path = out_path.with_suffix(".txt")
            _write_txt(out_path, req.title, req.q, answer, items, receipts)
            req_fmt = "txt"
        else:
            req_fmt = "docx"
    else:
        _write_txt(out_path, req.title, req.q, answer, items, receipts)
        req_fmt = "txt"

    size = 0
    try:
        size = out_path.stat().st_size
    except Exception:
        size = 0

    # Write JSON sidecar with receipts and hash
    sidecar = out_path.with_suffix(out_path.suffix + ".json")
    side = {
        "version": 1,
        "title": req.title,
        "created_at": _dt.datetime.utcnow().isoformat() + "Z",
        "path": str(out_path),
        "format": req_fmt,
        "bundle_hash": bundle_hash,
        "receipts": receipts,
        "citations": items,
    }
    try:
        sidecar.write_text(json.dumps(side, indent=2, sort_keys=True), encoding="utf-8")
    except Exception:
        pass

    # Auto-export proof artifacts: copy sidecar and summary into proof/RECEIPTS_SAMPLE/<timestamp>
    try:
        from datetime import datetime

        ts = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
        proof_dir = Path("proof") / "RECEIPTS_SAMPLE" / ts
        proof_dir.mkdir(parents=True, exist_ok=True)
        # Copy sidecar JSON into proof folder
        try:
            import shutil as _sh

            _sh.copy2(str(sidecar), str(proof_dir / sidecar.name))
        except Exception:
            # fallback: attempt to write a new file with same contents
            (proof_dir / sidecar.name).write_text(json.dumps(side, indent=2, sort_keys=True), encoding="utf-8")
        # Minimal summary
        (proof_dir / "summary.txt").write_text(
            f"bundle_hash={bundle_hash}\nreport={out_path}\n", encoding="utf-8"
        )
    except Exception:
        pass

    # If abstained, mark in response
    abstain = bool(qres.get("abstain"))
    reason = qres.get("reason") if abstain else None

    return {
        "path": str(out_path),
        "format": req_fmt,
        "size": int(size),
        "bundle_hash": bundle_hash,
        "receipt": qres.get("receipt"),
        "ingest_receipt": qres.get("ingest_receipt"),
        "settle_receipt": qres.get("settle_receipt"),
        "citations": items,
        "abstain": abstain or None,
        "reason": reason,
    }
